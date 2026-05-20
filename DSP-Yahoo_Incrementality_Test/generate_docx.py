#!/usr/bin/env python3
"""Generate Yahoo + DSP Geo Test design doc as .docx, following the Samsung CTV template."""

import os, pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(__file__)
mmt  = pd.read_csv(os.path.join(BASE, "mmt_group_assignments.csv"))
mmt["weekly_enrollments"] = (mmt["total_enrollments"] / (mmt["n_days"] / 7)).round(0).astype(int)

CHART = os.path.join(BASE, "mmt_balance_power.png")
OUT   = os.path.join(BASE, "Yahoo_DSP_Geo_Test_Design.docx")

YAHOO_BLUE    = RGBColor(0x15, 0x65, 0xC0)
DSP_ORANGE    = RGBColor(0xE6, 0x51, 0x00)
CTRL_GREEN    = RGBColor(0x2E, 0x7D, 0x32)
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
TABLE_HEADER  = RGBColor(0x1A, 0x23, 0x7E)

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "E0E4F0"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    p    = doc.add_heading(text, level=level)
    run  = p.runs[0]
    run.font.color.rgb = color or HEADING_COLOR
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(15)
    elif level == 2:
        run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_bullet(doc, text, bold_prefix=None, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after   = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.size = Pt(11)
    return p

# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading("Yahoo + DSP Display — Geo Incrementality Test Design", 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title.runs:
    run.font.color.rgb = YAHOO_BLUE
    run.font.size = Pt(18)
    run.font.bold = True
title.paragraph_format.space_after = Pt(4)

sub = doc.add_paragraph("Matched Market Test (MMT)  |  3-Cell Design  |  MDF-1629")
sub.paragraph_format.space_after = Pt(14)
for r in sub.runs:
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── What we're testing ─────────────────────────────────────────────────────────
add_heading(doc, "What We're Testing", level=1)
add_body(doc,
    "Whether running Yahoo display ads and DSP (Scope3) display ads in select U.S. markets "
    "drives more enrollments to Chime. The two partners are tested independently against the "
    "same control group, allowing us to measure the incremental lift of each channel separately."
)

# ── Test Design ────────────────────────────────────────────────────────────────
add_heading(doc, "Test Design", level=1)
add_bullet(doc, "10 Yahoo Test markets → Yahoo display ads ON")
add_bullet(doc, "10 DSP (Scope3) Test markets → DSP display ads ON")
add_bullet(doc, "10 Control markets → no display ads from either partner")
add_bullet(doc,
    "Enrollments are compared within each matched triplet to estimate incremental lift "
    "driven by each partner's display spend"
)

# ── Market Selection ───────────────────────────────────────────────────────────
add_heading(doc, "Market Selection", level=1)
add_bullet(doc, "Excluded the 26 largest U.S. DMAs (too dominant — would skew results) and the 50 smallest markets (too few enrollments to detect meaningful lift)")
add_bullet(doc, "Excluded 20 additional DMAs already running in active geo tests")
add_bullet(doc,
    "Matched each Yahoo market with a similar DSP market and a Control market into a \"triplet\" — "
    "matched on historical weekly enrollment volume and trend correlation (Feb 2025 – Apr 2026)"
)
add_bullet(doc, "Cell assignment within each triplet done to maximise enrollment balance across all three cells")

# ── Matched Markets table ──────────────────────────────────────────────────────
add_heading(doc, "Matched Markets", level=1)
p = doc.add_paragraph(
    "Each row is a matched triplet. Ads run only in Yahoo Test and DSP Test markets; "
    "Control markets receive no display spend from either partner."
)
p.paragraph_format.space_after = Pt(8)
for r in p.runs:
    r.font.size = Pt(11)

col_labels = ["Yahoo Test (ads ON)", "DSP (Scope3) Test (ads ON)", "Control Holdout (ads OFF)"]
col_colors = ["1565C0", "E65100", "2E7D32"]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

# Header row
hdr_cells = table.rows[0].cells
for i, (label, color) in enumerate(zip(col_labels, col_colors)):
    hdr_cells[i].text = label
    set_cell_bg(hdr_cells[i], color)
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size  = Pt(10)
    hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Data rows — one per triplet
yahoo_dmas = mmt[mmt["cell"] == "Yahoo Test"].sort_values("triplet")["dma_name"].tolist()
dsp_dmas   = mmt[mmt["cell"] == "DSP (Scope3) Test"].sort_values("triplet")["dma_name"].tolist()
ctrl_dmas  = mmt[mmt["cell"] == "Control Holdout"].sort_values("triplet")["dma_name"].tolist()

for i, (y, d, c) in enumerate(zip(yahoo_dmas, dsp_dmas, ctrl_dmas)):
    row_cells = table.add_row().cells
    bg = "F5F8FF" if i % 2 == 0 else "FFFFFF"
    for j, (name, bg_c) in enumerate(zip([y, d, c], [bg, bg, bg])):
        row_cells[j].text = name
        set_cell_bg(row_cells[j], bg_c)
        p = row_cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.size = Pt(10)
        row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Set column widths
for row in table.rows:
    for cell in row.cells:
        cell.width = Inches(2.0)

doc.add_paragraph()  # spacing

# ── Pre-Test Validation ────────────────────────────────────────────────────────
add_heading(doc, "Pre-Test Validation", level=1)
add_body(doc,
    "Test and control markets show strong alignment in weekly enrollments over the pre-period "
    "(Feb 2025 – Apr 2026). Each triplet was selected to maximise the week-to-week correlation "
    "between its three markets, giving confidence that any divergence during the campaign can be "
    "attributed to the ads. The chart below shows the enrollment profile across matched markets — "
    "overlapping lines confirm tight matching."
)

# Embed chart
doc.add_picture(CHART, width=Inches(6.2))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph(
    "Fig 1: Left — weekly enrollment trend by cell. Centre — DMA enrollment profile "
    "(overlapping lines = good match). Right — minimum detectable effect vs. test duration."
)
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_after = Pt(10)
for r in cap.runs:
    r.font.size  = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── Test Recommendation ────────────────────────────────────────────────────────
add_heading(doc, "Test Recommendation", level=1)
add_body(doc,
    "We recommend concentrating the $300K budget per partner into a 3-week flight. "
    "Concentrating spend into fewer weeks produces a stronger weekly signal, which is easier "
    "to detect against normal week-to-week market fluctuations. Spreading the same budget over "
    "4–6 weeks weakens the detectable signal and reduces our confidence in the result."
)

add_bullet(doc, "3 weeks",              bold_prefix="Recommended flight duration:")
add_bullet(doc, "$300K per partner",    bold_prefix="Budget:")
add_bullet(doc, "$400",                 bold_prefix="Assumed iCPE:")
add_bullet(doc, "~750 per partner",     bold_prefix="Expected incremental enrollments:")
add_bullet(doc, "~$10K per market per week",  bold_prefix="Spend pacing:")
add_bullet(doc, "~4%",                  bold_prefix="Lift this budget is expected to drive:")
add_bullet(doc, "~70% for both Yahoo and DSP",   bold_prefix="Statistical power:")
add_bullet(doc,
    "Week 1 allows the campaigns to exit the platform learning phase; "
    "Weeks 2–3 are the primary measurement window",
    bold_prefix="Flight structure:"
)

doc.add_paragraph()

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(8)
note.paragraph_format.left_indent  = Inches(0.0)
r = note.add_run(
    "Note: ")
r.bold = True
r.font.size = Pt(10)
r2 = note.add_run(
    "Expected incremental enrollments are the same for both partners ($300K ÷ $400 iCPE = 750) "
    "and are held equal by design. Actual results will depend on delivery, true iCPE, and "
    "whether ads reach incremental audiences in each market."
)
r2.font.size  = Pt(10)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Document saved: {OUT}")
